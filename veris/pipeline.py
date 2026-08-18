"""Ingestion pipeline: file in, knowledge graph out.

    upload → validate → extract text → freeze canonical text + hash
           → read declared provenance → parse structure into entities
           → verify every span → persist source, document, evidence, entities
           → retain the original artifact

Source-agnostic by construction: format handling is confined to `read_text`,
so adding a connector means adding one function, not touching the domain model.

Document structure survives ingestion (§13). Headings, section numbers and
provision titles become entity locators rather than being flattened away, which
is what makes a citation like "Medication Wasting Procedure §4 Document the
Waste" possible at all.
"""

from __future__ import annotations

import hashlib
import json
import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from .extract import TYPE_ROLE, extract, verify
from .store import Store

SUPPORTED = {".md", ".markdown", ".txt", ".pdf", ".docx"}
MAX_BYTES = 25 * 1024 * 1024


class IngestError(ValueError):
    """Raised for input the pipeline refuses. Messages are safe to surface."""


@dataclass
class IngestResult:
    document_id: str
    source_id: str
    title: str
    document_type: str
    entities: int
    rejected: int
    reused: bool = False


def read_text(path: Path) -> str:
    """Extract text, preserving structure. One function per format."""
    suffix = path.suffix.lower()
    if suffix in (".md", ".markdown", ".txt"):
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        import pypdf
        return "\n".join(p.extract_text() or "" for p in pypdf.PdfReader(str(path)).pages)
    if suffix == ".docx":
        import docx
        d = docx.Document(str(path))
        parts: list[str] = []
        for para in d.paragraphs:
            text = para.text.rstrip()
            if not text:
                continue
            # Preserve heading level as markdown so the section parser can see it.
            style = (para.style.name or "").lower()
            if style.startswith("heading"):
                level = "".join(ch for ch in style if ch.isdigit()) or "2"
                parts.append(f"{'#' * min(int(level), 6)} {text}")
            else:
                parts.append(text)
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    raise IngestError(f"Unsupported file type: {path.suffix}")


def parse_front_matter(text: str) -> tuple[dict, str]:
    """Flat key/value front matter. Not a YAML parser — the metadata here is
    flat, and a real parser would add a dependency for no gain. Unparseable
    front matter yields no metadata rather than an error, because metadata
    enriches a document and never grounds it."""
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 3)
    if end == -1:
        return {}, text
    meta: dict = {}
    for line in text[3:end].splitlines():
        if ":" not in line:
            continue
        k, _, v = line.partition(":")
        meta[k.strip()] = v.strip().strip('"').strip("'")
    return meta, text[end + 4:].lstrip("\n")


def validate(path: Path) -> None:
    if not path.is_file():
        raise IngestError(f"Not a file: {path.name}")
    if path.suffix.lower() not in SUPPORTED:
        raise IngestError(
            f"Unsupported file type {path.suffix!r}. Supported: "
            + ", ".join(sorted(SUPPORTED)))
    if path.stat().st_size > MAX_BYTES:
        raise IngestError(f"File exceeds {MAX_BYTES // (1024*1024)} MB limit")
    if path.stat().st_size == 0:
        raise IngestError("File is empty")


def _infer_type(meta: dict, name: str) -> str:
    if meta.get("document_type"):
        return meta["document_type"].upper()
    n = name.lower()
    for key, value in (("competenc", "COMPETENCY"), ("educat", "EDUCATION"),
                       ("orientation", "ORIENTATION"), ("procedure", "PROCEDURE"),
                       ("policy", "POLICY"), ("standard", "STANDARD")):
        if key in n:
            return value
    return "POLICY"


def ingest_file(store: Store, path: Path, data_dir: Path,
                default_source_type: str = "ORGANIZATIONAL") -> IngestResult:
    path = Path(path)
    validate(path)

    raw = read_text(path)
    meta, body = parse_front_matter(raw)
    if not body.strip():
        raise IngestError(f"No extractable text in {path.name}")

    # Canonical text is frozen here and never rewritten. Every span cites into
    # this file, so if the source changes its hash changes and every entity
    # citing it is stale by construction.
    text_sha = hashlib.sha256(body.encode("utf-8")).hexdigest()
    canonical_dir = data_dir / "canonical"
    canonical_dir.mkdir(parents=True, exist_ok=True)
    canonical_path = canonical_dir / f"{text_sha[:16]}.txt"
    canonical_path.write_text(body, encoding="utf-8")

    existing = store.one("SELECT * FROM documents WHERE text_sha256 = ?", (text_sha,))
    if existing:
        return IngestResult(existing["id"], existing["source_id"], existing["title"],
                            existing["document_type"], 0, 0, reused=True)

    doc_type = _infer_type(meta, path.name)
    title = meta.get("title") or re.sub(r"[_-]+", " ", path.stem)
    version = meta.get("version", "")

    # Originals are retained; provenance means being able to return to the artifact.
    originals = data_dir / "originals"
    originals.mkdir(parents=True, exist_ok=True)
    stored = originals / f"{text_sha[:16]}{path.suffix.lower()}"
    shutil.copy2(path, stored)

    source_id = store.add_source(
        source_type=meta.get("source_type", default_source_type).upper(),
        title=meta.get("source_title") or title,
        publisher=meta.get("publisher"),
        authority=meta.get("authority"),
        jurisdiction=meta.get("jurisdiction"),
        publication_date=meta.get("publication_date"),
        effective_date=meta.get("effective_date"),
        version=version,
        retrieval_date=meta.get("retrieval_date"),
        reference_url=meta.get("reference_url"),
        content_hash=text_sha,
    )
    document_id = store.add_document(
        source_id=source_id, document_type=doc_type, title=title, version=version,
        effective_date=meta.get("effective_date"), status=meta.get("status", "ACTIVE"),
        owner=meta.get("owner"), department=meta.get("department"),
        text_sha256=text_sha, char_count=len(body),
        storage_path=str(stored), canonical_path=str(canonical_path),
        metadata=json.dumps(meta),
    )

    entity_type, role = TYPE_ROLE.get(doc_type, ("POLICY_STATEMENT", "COMMITS"))
    standard_id = meta.get("standard") or title.split("—")[0].strip()
    found = extract(body, doc_type, standard_id if doc_type == "STANDARD" else title)
    ok, rejected = verify(found, body)

    for e in ok:
        ev_id = store.add_evidence(document_id, e.char_start, e.char_end,
                                   e.statement, e.title)
        store.add_entity(
            document_id=document_id, evidence_id=ev_id, entity_type=entity_type,
            role=role, locator=e.locator, statement=e.statement,
            subject=e.title, expects_document=e.expects_document,
            crosswalk=e.crosswalk, provenance_class="SOURCE_FACT",
        )
    store.log("document_ingested",
              f"{title} v{version or '-'} · {doc_type} · {len(ok)} entities")
    store.commit()
    return IngestResult(document_id, source_id, title, doc_type, len(ok), len(rejected))


def ingest_directory(store: Store, directory: Path, data_dir: Path,
                     default_source_type: str = "ORGANIZATIONAL") -> list[IngestResult]:
    results = []
    for path in sorted(Path(directory).rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED and not path.name.startswith("_"):
            try:
                results.append(ingest_file(store, path, data_dir, default_source_type))
            except IngestError as e:
                store.log("ingest_rejected", f"{path.name}: {e}")
    return results
